"""
Document Extractor

Supported formats:
- PDF (PyPDF2 + OCR fallback)
- Word (python-docx)
- Excel (openpyxl)
- Images (PaddleOCR)
- Plain text (UTF-8/GBK/Latin-1)
"""

import logging
from dataclasses import dataclass
from pathlib import Path
from typing import Optional, Union

logger = logging.getLogger(__name__)


@dataclass
class DocumentContent:
    """Container for extracted document content and metadata."""

    text: str  # Extracted text content
    metadata: dict  # Document metadata (pages, author, etc.)
    method: str  # Extraction method used (e.g., pypdf2, ocr)
    file_type: str  # Original file type (pdf, word, excel, etc.)


class DocumentExtractor:
    """
    Multi-format document text extractor.

    Supports PDF, Word, Excel, plain text, and image files with optional OCR fallback.
    """

    SUPPORTED_EXTENSIONS = {
        ".pdf": "pdf",
        ".docx": "word",
        ".doc": "word",
        ".xlsx": "excel",
        ".xls": "excel",
        ".txt": "text",
        ".md": "text",
        ".png": "image",
        ".jpg": "image",
        ".jpeg": "image",
        ".bmp": "image",
        ".tiff": "image",
    }

    def __init__(self, use_ocr: bool = True):
        """
        Initialize the document extractor.

        Args:
            use_ocr: Whether to enable OCR for scanned PDFs and images
        """
        self.use_ocr = use_ocr
        self.ocr_processor = None

        if use_ocr:
            try:
                from backend.services.document_processing.ocr_processor import (
                    OCRProcessor,
                )

                self.ocr_processor = OCRProcessor()
            except Exception as e:
                logger.warning(f"OCR processor initialization failed: {e}")

    def extract(
        self, file_path: Union[str, Path], progress_callback=None
    ) -> DocumentContent:
        """
        Extract text content from a document file.

        Args:
            file_path: Path to the document file
            progress_callback: Optional callback(stage, status, progress, detail)
                for reporting extraction progress to SSE clients.

        Returns:
            DocumentContent with extracted text and metadata

        Raises:
            ValueError: If the file type is unsupported
            FileNotFoundError: If the file does not exist
        """
        file_path = Path(file_path)

        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        file_ext = file_path.suffix.lower()
        if file_ext not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(f"Unsupported file extension: {file_ext}")

        file_type = self.SUPPORTED_EXTENSIONS[file_ext]

        # Route to the appropriate extraction method
        if file_type == "pdf":
            return self._extract_pdf(file_path, progress_callback=progress_callback)
        elif file_type == "word":
            return self._extract_word(file_path)
        elif file_type == "excel":
            return self._extract_excel(file_path)
        elif file_type == "text":
            return self._extract_text(file_path)
        elif file_type == "image":
            return self._extract_image(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")

    def _extract_pdf(self, file_path: Path, progress_callback=None) -> DocumentContent:
        """Extract text from a PDF file, with OCR fallback for scanned pages."""
        try:
            import fitz  # PyMuPDF

            doc = fitz.open(file_path)
            total_pages = len(doc)
            text_parts = []
            ocr_pages = []

            for page_num in range(total_pages):
                page = doc[page_num]
                page_text = page.get_text().strip()

                if page_text:
                    text_parts.append(page_text)
                elif self.use_ocr and self.ocr_processor:
                    # Scanned page — render to image and OCR
                    logger.info(
                        f"PDF page {page_num + 1} has no extractable text, "
                        "falling back to OCR"
                    )
                    if progress_callback:
                        progress_callback(
                            "ocr", "running",
                            (page_num + 1) / total_pages,
                            f"OCR: page {page_num + 1}/{total_pages}",
                        )
                    try:
                        pix = page.get_pixmap(dpi=200)
                        import tempfile

                        with tempfile.NamedTemporaryFile(
                            suffix=".png", delete=False
                        ) as tmp:
                            pix.save(tmp.name)
                            ocr_result = self.ocr_processor.process(tmp.name)
                            if ocr_result.text.strip():
                                text_parts.append(ocr_result.text)
                                ocr_pages.append(page_num + 1)
                            Path(tmp.name).unlink(missing_ok=True)
                    except Exception as ocr_err:
                        logger.warning(
                            f"OCR fallback failed for page {page_num + 1}: {ocr_err}"
                        )

                # Report extraction progress
                if progress_callback:
                    progress_callback(
                        "extract", "running",
                        (page_num + 1) / total_pages,
                        f"Extracting text: {page_num + 1}/{total_pages} pages",
                    )

            full_text = "\n\n".join(text_parts)

            pdf_metadata = doc.metadata or {}
            metadata = {
                "num_pages": total_pages,
                "author": pdf_metadata.get("author", "Unknown"),
                "title": pdf_metadata.get("title", "Unknown"),
            }
            if ocr_pages:
                metadata["ocr_pages"] = ocr_pages
                if progress_callback:
                    progress_callback(
                        "ocr", "completed", 1.0,
                        f"OCR completed ({len(ocr_pages)} pages)",
                    )

            method = "pymupdf+ocr" if ocr_pages else "pymupdf"

            doc.close()

            return DocumentContent(
                text=full_text,
                metadata=metadata,
                method=method,
                file_type="pdf",
            )

        except Exception as e:
            logger.error(f"Failed to extract PDF content: {e}")
            raise

    def _extract_word(self, file_path: Path) -> DocumentContent:
        """Extract text from a Word document."""
        try:
            from docx import Document

            doc = Document(file_path)

            # Extract paragraphs
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

            # Extract table content
            tables_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells)
                    tables_text.append(row_text)

            # Combine paragraphs and tables
            all_text = "\n".join(paragraphs)
            if tables_text:
                all_text += "\n\n" + "\n".join(tables_text)

            metadata = {
                "num_paragraphs": len(paragraphs),
                "num_tables": len(doc.tables),
            }

            return DocumentContent(
                text=all_text,
                metadata=metadata,
                method="python-docx",
                file_type="word",
            )

        except Exception as e:
            logger.error(f"Failed to extract Word document content: {e}")
            raise

    def _extract_excel(self, file_path: Path) -> DocumentContent:
        """Extract text from an Excel spreadsheet."""
        try:
            import openpyxl

            workbook = openpyxl.load_workbook(file_path, data_only=True)

            sheets_text = []
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]

                # Add sheet header
                sheet_text = [f"### {sheet_name} ###"]

                for row in sheet.iter_rows(values_only=True):
                    row_values = [str(cell) if cell is not None else "" for cell in row]
                    if any(row_values):  # Skip empty rows
                        sheet_text.append(" | ".join(row_values))

                sheets_text.append("\n".join(sheet_text))

            full_text = "\n\n".join(sheets_text)

            metadata = {
                "num_sheets": len(workbook.sheetnames),
                "sheet_names": workbook.sheetnames,
            }

            return DocumentContent(
                text=full_text,
                metadata=metadata,
                method="openpyxl",
                file_type="excel",
            )

        except Exception as e:
            logger.error(f"Failed to extract Excel content: {e}")
            raise

    def _extract_text(self, file_path: Path) -> DocumentContent:
        """Extract content from a plain text file."""
        try:
            # Try multiple encodings
            encodings = ["utf-8", "gbk", "gb2312", "latin-1"]

            for encoding in encodings:
                try:
                    with open(file_path, "r", encoding=encoding) as f:
                        text = f.read()

                    metadata = {
                        "encoding": encoding,
                        "num_lines": len(text.splitlines()),
                        "num_chars": len(text),
                    }

                    return DocumentContent(
                        text=text,
                        metadata=metadata,
                        method=f"text-{encoding}",
                        file_type="text",
                    )

                except UnicodeDecodeError:
                    continue

            raise ValueError("Unable to decode text file with any supported encoding")

        except Exception as e:
            logger.error(f"Failed to extract text file content: {e}")
            raise

    def _extract_image(self, file_path: Path) -> DocumentContent:
        """Extract text from an image file using OCR."""
        if not self.use_ocr or not self.ocr_processor:
            raise ValueError("OCR is not available for image extraction")

        try:
            from PIL import Image

            # Get image dimensions and mode
            with Image.open(file_path) as img:
                width, height = img.size
                mode = img.mode

            # Run OCR on the image
            ocr_result = self.ocr_processor.process(file_path)

            metadata = {
                "width": width,
                "height": height,
                "mode": mode,
                "ocr_confidence": ocr_result.confidence,
                "ocr_method": ocr_result.method,
            }

            return DocumentContent(
                text=ocr_result.text,
                metadata=metadata,
                method="ocr",
                file_type="image",
            )

        except Exception as e:
            logger.error(f"Image OCR extraction failed: {e}")
            raise


# Convenience function
def process_document(
    file_path: Union[str, Path],
    use_ocr: bool = True,
) -> DocumentContent:
    """
    Extract text from a document file (convenience wrapper).

    Args:
        file_path: Path to the document file
        use_ocr: Whether to enable OCR for images and scanned PDFs

    Returns:
        DocumentContent with extracted text and metadata
    """
    extractor = DocumentExtractor(use_ocr=use_ocr)
    return extractor.extract(file_path)
